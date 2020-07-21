from docx import Document
import re
import os
import tkinter as tk
from tkinter import filedialog, Text

root = tk.Tk()
dokos = []

canvas = tk.Canvas(root, height = 150, width = 200, bg='#5BCF8C')
canvas.pack()
frame = tk.Frame(root, bg='#B5E771')
frame.place(relwidth=0.8, relheight=0.8, relx=0.1, rely=0.1)
#def addDocs():
#    filename = filedialog.askopenfilename(initialdir="", title="Select File", filetypes=(("Documents", "*.docx"),))
#    dokos.append(filename)
#    print(filename)
#    for docs in dokos:
#        label = tk.Label(frame, text=docs)
#        label.pack()
def runIt():
    allDiagn = Document()
    for files in os.listdir():
        if files.endswith('docx'):
            if files != 'All The Diagnoses.docx':
                doc = Document(files)
                diagn = ''
                for num in range (4):
                    diagn = doc.paragraphs[num].text
                    if 'diagnoses' in diagn:
                        break
                sub = re.search('diagnoses of (.*?)\.', diagn)
                if sub is not None:
                    subStr = sub.group(1)
                    ar = re.split(',', subStr)
                    allDiagn.add_paragraph(files.replace(".docx", ":") + '\n' + '|'.join(ar))
    allDiagn.save("All The Diagnoses.docx")
    os.startfile("All The Diagnoses.docx")
label = tk.Label(frame, text="I LOVE YOUUU MWUAH~")
label.pack()
runAll = tk.Button(root, text='Run For All Documents', command=runIt)
runAll.pack()
#openSome = tk.Button(root, text='Select Some Files', command=addDocs)
#openSome.pack()

root.mainloop()