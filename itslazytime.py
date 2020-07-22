from docx import Document
import re
import os
import tkinter as tk
from tkinter import filedialog, Text

root = tk.Tk()
root.title('Failure')
root.geometry("300x400")
root.configure(background='#f5776e')
dokos = []

listbox = tk.Listbox(root)
listbox.pack(pady=15)

def addDocs():
    filename = filedialog.askopenfilename(initialdir="", title="Select File", filetypes=(("Documents", "*.docx"),))
    dokos.append(filename)
    #print(filename)
    listbox.insert('end', filename[filename.rfind('/')+1:])

def removeItem():
    dokos.pop(listbox.get(0, 'end').index(listbox.get('active')))
    #print(dokos)
    listbox.delete('anchor')
    

def runSelected():
    allDiagn = Document()
    for files in dokos:
        if files.endswith('docx'):
            doc = Document(files)
            diagn = ''
            for num in range (4):
                diagn = doc.paragraphs[num].text
                if 'diagnoses' in diagn:
                    break
            sub = re.search('diagnoses of (.*?)\.', diagn)
            if sub is not None:
                subStr = sub.group(1)
                ar = re.split(',', subStr)
                allDiagn.add_paragraph(files[files.rfind('/')+1:].replace(".docx", ":") + '\n' + '|'.join(ar))
    allDiagn.save("All The Diagnoses.docx")
    os.startfile("All The Diagnoses.docx")

def runIt():
    allDiagn = Document()
    for files in os.listdir():
        if files.endswith('docx'):
            if files != 'All The Diagnoses.docx':
                doc = Document(files)
                diagn = ''
                for num in range (4):
                    diagn = doc.paragraphs[num].text
                    if 'diagnoses' in diagn:
                        break
                sub = re.search('diagnoses of (.*?)\.', diagn)
                if sub is not None:
                    subStr = sub.group(1)
                    ar = re.split(',', subStr)
                    allDiagn.add_paragraph(files.replace(".docx", ":") + '\n' + '|'.join(ar))
    allDiagn.save("All The Diagnoses.docx")
    os.startfile("All The Diagnoses.docx")
#label = tk.Label(frame, text="I LOVE YOUUU MWUAH~")
#label.pack()

openSome = tk.Button(root, text='Select Files', command=addDocs)
openSome.pack(pady=5)

rmList = tk.Button(root, text='Remove from list', command=removeItem)
rmList.pack(pady=5)

runSelect = tk.Button(root, text='Run Selected', command=runSelected)
runSelect.pack(pady=5)

#runAll = tk.Button(root, text='Run For All Documents', command=runIt)
#runAll.pack(pady=5)



root.mainloop()